#!/usr/bin/env python3
"""
Создание демонстрационных DOCX файлов с протоколами
"""

import sys
from pathlib import Path

# Добавляем корневую директорию проекта в Python path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Inches

def create_minutes_01():
    """Создание minutes_01_discussion.docx"""
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Переписка 1: Обсуждение сроков производства', 0)
    title.alignment = 1  # Центрирование
    
    # Добавляем пустую строку
    doc.add_paragraph('')
    
    # Алексей
    alexey_para = doc.add_paragraph()
    alexey_para.add_run('Алексей:').bold = True
    doc.add_paragraph(
        'По заказу A101 есть задержка с материалом от поставщика. Возможен перенос начала с 1 июля на 5 июля. '
        'Как думаешь, успеем к 10 июля?'
    )
    
    # Пустая строка
    doc.add_paragraph('')
    
    # Марина
    marina_para = doc.add_paragraph()
    marina_para.add_run('Марина:').bold = True
    doc.add_paragraph(
        'Алексей, если увеличим скорость линии до 120 ед./день, то успеем. '
        'Но нужно согласовать с руководством.'
    )
    
    # Пустая строка  
    doc.add_paragraph('')
    
    # Алексей ответ
    alexey2_para = doc.add_paragraph()
    alexey2_para.add_run('Алексей:').bold = True
    doc.add_paragraph('Хорошо, я уточню. Дам ответ завтра.')
    
    # Сохраняем файл
    output_path = project_root / 'demo' / 'minutes' / 'minutes_01_discussion.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(str(output_path))
    print(f"✅ Создан файл: {output_path}")
    return output_path

def create_minutes_02():
    """Создание minutes_02_planned_downtime.docx"""
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Переписка 2: Утверждение планового простоя', 0)
    title.alignment = 1  # Центрирование
    
    # Добавляем пустую строку
    doc.add_paragraph('')
    
    # Сергей
    sergey_para = doc.add_paragraph()
    sergey_para.add_run('Сергей:').bold = True
    
    # Основной текст с важной информацией
    main_text = doc.add_paragraph(
        'Коллеги, утверждена остановка линии с 1 по 31 августа на плановое обслуживание. '
        'Все работы должны быть завершены до сентября.'
    )
    
    # Пустая строка
    doc.add_paragraph('')
    
    # Выделенная цитата (как указано в требованиях)
    quote_para = doc.add_paragraph()
    quote_run = quote_para.add_run(
        '"Утверждена остановка линии с 1 по 31 августа на плановое обслуживание."'
    )
    quote_run.italic = True
    
    # Сохраняем файл
    output_path = project_root / 'demo' / 'minutes' / 'minutes_02_planned_downtime.docx'
    
    doc.save(str(output_path))
    print(f"✅ Создан файл: {output_path}")
    return output_path

def create_minutes_03():
    """Создание дополнительного файла с упоминанием линии и синонимов"""
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Переписка 3: Техническое совещание по линии', 0)
    title.alignment = 1  # Центрирование
    
    # Добавляем пустую строку
    doc.add_paragraph('')
    
    # Игорь
    igor_para = doc.add_paragraph()
    igor_para.add_run('Игорь:').bold = True
    
    # Текст с упоминанием линии и синонимов
    doc.add_paragraph(
        'Добрый день! По 66-й линии нужно провести техническое обслуживание системы freeze-dry. '
        'Планируем остановить freeze-dry линию на 3 дня в сентябре для замены компонентов.'
    )
    
    # Пустая строка
    doc.add_paragraph('')
    
    # Анна
    anna_para = doc.add_paragraph()
    anna_para.add_run('Анна:').bold = True
    doc.add_paragraph(
        'Игорь, согласовано. Линия_66 будет остановлена с 15 по 17 сентября. '
        'Все операции по фриз-драй оборудованию выполним в указанные сроки.'
    )
    
    # Сохраняем файл
    output_path = project_root / 'demo' / 'minutes' / 'minutes_03_line_maintenance.docx'
    
    doc.save(str(output_path))
    print(f"✅ Создан файл: {output_path}")
    return output_path

if __name__ == "__main__":
    try:
        created_files = []
        
        # Создаем все файлы
        created_files.append(create_minutes_01())
        created_files.append(create_minutes_02())
        created_files.append(create_minutes_03())
        
        print(f"\n✅ Созданы файлы протоколов:")
        for file_path in created_files:
            print(f"   - {file_path}")
            
    except Exception as e:
        print(f"❌ Ошибка создания файлов: {e}")
        sys.exit(1)